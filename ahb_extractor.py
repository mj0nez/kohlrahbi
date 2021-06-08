from pathlib import Path
from typing import List

import docx
import pandas as pd
from docx.table import Table
from docx.text.paragraph import Paragraph

from check_row_type import RowType, define_row_type
from get_sections import get_chapter_with_ahb_tables, iter_block_items
from write_functions import write_new_row_in_dataframe

directory_path = Path.cwd() / "documents"
file_name = "UTILMD_AHB_WiM_3_1c_2021_04_01_2021_03_30.docx"

path_to_file = directory_path / file_name


def get_tabstop_positions(paragraph):
    tabstop_positions: List = []
    for tabstop in paragraph.paragraph_format.tab_stops._pPr.tabs:
        tabstop_positions.append(tabstop.pos)
    return tabstop_positions


def main():
    try:
        doc = docx.Document(path_to_file)  # Creating word reader object.

        # TODO for each section get header to get prüfidentifaktoren for dataframe header

        # get_chapter_with_ahb_tables(document=doc)
        for item in iter_block_items(parent=doc):
            if isinstance(item, Paragraph) and "Heading" in item.style.name:
                chapter_title = item.text
            elif isinstance(item, Table) and item.cell(row_idx=0, col_idx=0).text == "EDIFACT Struktur":
                print(chapter_title)
                print([cell.text for cell in item.row_cells(0)])
                print("\n NEW TABLE \n")

                header_cells = [cell.text for cell in item.row_cells(0)]
                look_up_term = "Prüfidentifikator"
                cutter_index = header_cells[-1].find(look_up_term) + 1
                # +1 cause of \t after Prüfidentifikator
                pruefidentifikatoren: List = header_cells[-1][cutter_index + len(look_up_term) :].split("\t")

                tabstop_positions: List = get_tabstop_positions(item.cell(row_idx=2, col_idx=1).paragraphs[0])

                pass

                # read the following information from the table
                # * all Prüfidentifikatoren
                # * tabstop position in middle cell
                # initial new dataframe with the known Prüfidentifikatoren
            # elif isinstance(item, Table):
            #     print(item.cell(0, 0).text)
            else:
                continue

        # Initialize help variables
        last_two_row_types: List = [RowType.EMPTY, RowType.EMPTY]
        actual_df_row_index: int = 0

        df = pd.DataFrame(
            columns=[
                "Segment Gruppe",
                "Segment",
                "Datenelement",
                "Codes und Qualifier",
                "Beschreibung",
                "11039",  # should be taken from table header
                "11040",  # should be taken from table header
                "11041",  # should be taken from table header
                "Bedingung",
            ],
            dtype="str",
        )

        # for table in doc.tables[1 : len(doc.tables)]:
        for table in doc.tables[1:20]:

            for row in range(len(table.rows)):

                index_for_middle_column = 1
                # initial empty list for the next row in the dataframe
                actual_dataframe_row = (len(df.columns)) * [""]

                # idea: cause of double information in 4 column tables, remove entry [2] from row_cell_texts_as_list
                row_cell_texts_as_list = [cell.text for cell in table.row_cells(row)]
                if table._column_count == 4:
                    # remove redundant information for tables with 4 columns
                    if (
                        row_cell_texts_as_list[0] == row_cell_texts_as_list[1]
                        and row_cell_texts_as_list[2] == row_cell_texts_as_list[3]
                    ):
                        # HEADER looks like
                        # 0:'EDIFACT Struktur'
                        # 1:'EDIFACT Struktur'
                        # 2:'Beschreibung\tKündigung\tBestätigung\tAblehnung\tBedingung\n\tMSB \tKündigung\tKündigung\n\tMSB \tMSB \nKommunikation von\tMSBN an\tMSBA an\tMSBA an\n\tMSBA\tMSBN\tMSBN\nPrüfidentifikator\t11039\t11040\t11041'
                        # 3:'Beschreibung\tKündigung\tBestätigung\tAblehnung\tBedingung\n\tMSB \tKündigung\tKündigung\n\tMSB \tMSB \nKommunikation von\tMSBN an\tMSBA an\tMSBA an\n\tMSBA\tMSBN\tMSBN\nPrüfidentifikator\t11039\t11040\t11041'
                        # len():4
                        del row_cell_texts_as_list[1]
                        row_cell_texts_as_list[2] = ""
                    elif row_cell_texts_as_list[1] == row_cell_texts_as_list[2]:
                        # Dataelement row with header in the table
                        # 0:'SG2\tNAD\t3035'
                        # 1:'SG2\tNAD\t3035'
                        # 2:'MR\tNachrichtenempfänger\tX\tX\tX'
                        # 3:''
                        # len():4
                        del row_cell_texts_as_list[1]
                    elif row_cell_texts_as_list[0] == row_cell_texts_as_list[1]:
                        del row_cell_texts_as_list[1]

                    index_for_middle_column = 2

                actual_edifact_struktur_cell = table.row_cells(row)[0]

                # check for row type
                actual_row_type = define_row_type(
                    table=table,
                    edifact_struktur_cell=actual_edifact_struktur_cell,
                    text_in_row_as_list=row_cell_texts_as_list,
                )
                print(actual_row_type.name)

                # write actual row into dataframe

                #
                if actual_row_type is RowType.EMPTY and last_two_row_types[0] is RowType.HEADER:
                    actual_df_row_index = actual_df_row_index - 1
                    actual_df_row_index = write_new_row_in_dataframe(
                        row_type=last_two_row_types[1],
                        table=table,
                        row=row,
                        index_for_middle_column=index_for_middle_column,
                        dataframe=df,
                        dataframe_row_index=actual_df_row_index,
                        dataframe_row=actual_dataframe_row,
                        row_cell_texts_as_list=row_cell_texts_as_list,
                    )

                else:
                    actual_df_row_index = write_new_row_in_dataframe(
                        row_type=actual_row_type,
                        table=table,
                        row=row,
                        index_for_middle_column=index_for_middle_column,
                        dataframe=df,
                        dataframe_row_index=actual_df_row_index,
                        dataframe_row=actual_dataframe_row,
                        row_cell_texts_as_list=row_cell_texts_as_list,
                    )

                # remember last row type for empty cells
                last_two_row_types[1] = last_two_row_types[0]
                last_two_row_types[0] = actual_row_type

        df.to_csv("export.csv")
        # df.to_excel("export.xlsx")
        print(len(doc.tables[1].columns))

    except IOError:
        print("There was an error opening the file!")
        return


if __name__ == "__main__":
    main()
