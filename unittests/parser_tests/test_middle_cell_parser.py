import pytest
from kohlrahbi.parser import middle_cell_parser

import docx
from docx.shared import Inches


class TestMiddleCellParser:
    """
    This test class collects all tests in context of the middle cell parser
    """

    @pytest.mark.parametrize(
        ["left_indent_position", "expected_result"],
        [
            pytest.param(Inches(1), True, id="Paragraph contains code"),
            pytest.param(Inches(4), False, id="Paragraph does NOT contains code"),
        ],
    )
    def test_does_paragraph_contain_qualifier_or_code(self, left_indent_position, expected_result):
        doc = docx.Document()
        test_paragraph = doc.add_paragraph("\tMuss\tMuss\tMuss")
        test_paragraph.paragraph_format.left_indent = Inches(1)
        # though it feels weird to use Inches, stick to it! I tried to use the length class Mm()
        # but there was always a difference between:
        #   absolute_length = Mm(1)
        #   test_paragraph.paragraph_format.left_indent = Mm(1)
        #   absolute_length != test_paragraph.paragraph_format.left_indent
        # but I don't know why ...

        assert (
            middle_cell_parser.does_paragraph_contain_qualifier_or_code(
                paragraph=test_paragraph, left_indent_position=left_indent_position
            )
            is expected_result
        )

    def test_has_paragraph_tabstops(self):
        doc = docx.Document()
        test_paragraph = doc.add_paragraph("\tMuss\tMuss\tMuss")
        test_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))

        assert middle_cell_parser.has_paragraph_tabstops(paragraph=test_paragraph)
